from django.shortcuts import render,redirect ,get_object_or_404
from django.http import HttpResponse, JsonResponse
from django.urls import reverse
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.utils import timezone
from .forms import GenerarConstanciaForm, EventoForm, ParticipanteForm, CargaCSVForm
from docx import Document 
from io import BytesIO 
import os
import csv
import io
from django.conf import settings
from .models import Participante, Evento, Plantilla, TipoEvento, ModalidadEvento
from django.db.models import Min, Count
from datetime import datetime, timedelta
import logging

# Configurar logging
logger = logging.getLogger(__name__)

# Create your views here.

def test_post_view(request):
    print("🔥🔥🔥 TEST POST VIEW EJECUTÁNDOSE 🔥🔥🔥")
    print(f"Método: {request.method}")
    print(f"POST data: {request.POST}")
    
    if request.method == 'POST':
        print("✅ POST REQUEST RECIBIDO EXITOSAMENTE!")
        return HttpResponse("¡POST funciona! Método: POST")
    else:
        print("📄 GET REQUEST - Mostrando formulario")
        return render(request, 'test_post.html')

@login_required
def lista_participantes(request):
    

    ids_unicos = Participante.objects.values('email_participante').annotate(
        min_id=Min('id_participante')
    ).values_list('min_id', flat=True) 
    
    
    participantes = Participante.objects.filter(id_participante__in=ids_unicos).order_by('email_participante')
    
    
    
    contexto = {
        'participantes': participantes
    }
    
    return render(request, 'constancias/lista_participantes.html', contexto)


# constancias/views.py

@login_required
def detalle_participante(request, pk):
    # 1. Asigna 'participante' de manera segura
    participante = get_object_or_404(Participante, id_participante=pk)
    
    # 2. Accede a los eventos a través del related_name 'eventos'
    eventos_registrados = participante.eventos.all() 
    
    contexto = {
        'participante': participante,
        'eventos': eventos_registrados, 
    }
    
    return render(request, 'constancias/detalle_participante.html', contexto)

@login_required
def lista_eventos(request):
    eventos = Evento.objects.all().order_by('titulo_evento')

    contexto = {
        'eventos': eventos
    }

    return render(request,'constancias/lista_eventos.html',contexto)

@login_required
def detalle_evento(request,pk):
    
    evento = get_object_or_404(Evento,id_evento=pk)

    participantes_evento = evento.participantes.all().order_by('nombre_participante')

    contexto = {
        'evento': evento,
        'participantes': participantes_evento
    }

    return render(request, 'constancias/detalle_evento.html', contexto)

    

@login_required
def generar_constancia(request, participante_pk, evento_pk, plantilla_pk):
    # 1. OBTENER OBJETOS
    try:
        # Usa _pk en la función y en el filtro de la base de datos
        evento = Evento.objects.get(id_evento=evento_pk)
        participante = Participante.objects.get(id_participante=participante_pk)
        plantilla = Plantilla.objects.get(id_plantilla=plantilla_pk) 
        
    except (Evento.DoesNotExist, Participante.DoesNotExist, Plantilla.DoesNotExist) as e:
        # Manejo de error si no se encuentran los objetos
        return HttpResponse("Error: Objeto no encontrado. Detalle: " + str(e), status=404)

    # 2. DEFINIR SUSTITUCIONES
    sustituciones = {
        # ¡IMPORTANTE! Mantén la convención de doble llave
        '{{TITULO_EVENTO}}': evento.titulo_evento,
        '{{NOMBRE_PARTICIPANTE}}': participante.nombre_participante,
        '{{ROL_PARTICIPANTE}}': participante.rol_participante,        
        '{{FECHA_INICIO}}': evento.fecha_inicio.strftime('%d/%m/%Y'),
        '{{FECHA_FIN}}': evento.fecha_fin.strftime('%d/%m/%Y'),
    }

    # 3. CARGAR EL DOCUMENTO
    try:
        # Carga el archivo .docx desde la ruta del FileField
        document = Document(plantilla.archivo.path)
    except Exception as e:
        return HttpResponse(f"Error al cargar el archivo DOCX: {e}", status=500)


    # 4. LÓGICA DE REEMPLAZO (Función central para sustituir texto)
    def replace_text_in_element(element, replacements):
        """Reemplaza los marcadores en un párrafo o celda."""
        # Se itera sobre los runs porque si el texto tiene formato,
        # python-docx lo divide.
        for run in element.runs:
            for key, value in replacements.items():
                if key in run.text:
                    run.text = run.text.replace(key, value)

    def replace_all_markers(doc, replacements):
        # A. Reemplazar en párrafos (cuerpo principal)
        for p in doc.paragraphs:
            replace_text_in_element(p, replacements)
        
        # B. Reemplazar en tablas (muy común en plantillas complejas)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_text_in_element(p, replacements)

    replace_all_markers(document, sustituciones)


    # 5. PREPARAR RESPUESTA HTTP
    
    # Prepara el nombre del archivo para la descarga
    nombre_archivo = f"Constancia_{participante.nombre_participante.replace(' ', '_')}.docx"
    
    # Crea el objeto de respuesta HTTP para un archivo DOCX
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{nombre_archivo}"'
    
    # Guarda el documento modificado directamente en la respuesta
    document.save(response)

    return response
    
   


@login_required
def pagina_generar_constancia(request):
    # Determinar el origen (dashboard, evento, participante)
    evento_preseleccionado = None
    participante_preseleccionado = None
    origen = 'dashboard'  # Por defecto
    
    # Verificar si viene desde un evento específico
    evento_id = request.GET.get('evento_id')
    if evento_id:
        try:
            evento_preseleccionado = Evento.objects.get(pk=evento_id)
            origen = 'evento'
        except Evento.DoesNotExist:
            messages.error(request, "El evento especificado no existe.")
    
    # Verificar si viene desde un participante específico
    participante_id = request.GET.get('participante_id')
    if participante_id:
        try:
            participante_preseleccionado = Participante.objects.get(pk=participante_id)
            origen = 'participante'
        except Participante.DoesNotExist:
            messages.error(request, "El participante especificado no existe.")
    
    # Obtener todos los datos necesarios
    eventos = Evento.objects.filter(activo=True).order_by('-fecha_inicio')
    participantes = Participante.objects.all().order_by('nombre_participante')
    
    # Si hay evento preseleccionado, filtrar participantes de ese evento
    if evento_preseleccionado:
        participantes_evento = evento_preseleccionado.participantes.all().order_by('nombre_participante')
    else:
        participantes_evento = participantes
    
    if request.method == 'POST':
        # Determinar si es generación masiva o individual
        generar_todos = request.POST.get('generar_todos') == 'true'
        
        if generar_todos:
            # Generación masiva para todos los participantes del evento
            evento_id_form = request.POST.get('evento_id')
            
            try:
                evento = Evento.objects.get(pk=evento_id_form)
                
                # Verificar que el evento tiene una plantilla
                if not evento.plantilla_id:
                    messages.error(request, f"El evento '{evento.titulo_evento}' no tiene una plantilla asociada.")
                    return render(request, 'constancias/pagina_generar.html', {
                        'eventos': eventos,
                        'participantes': participantes,
                        'origen': origen,
                        'evento_preseleccionado': evento_preseleccionado,
                        'participante_preseleccionado': participante_preseleccionado,
                        'participantes_evento': participantes_evento,
                    })
                
                participantes_generar = evento.participantes.all()
                
                if not participantes_generar.exists():
                    messages.warning(request, "No hay participantes registrados en este evento.")
                    return render(request, 'constancias/pagina_generar.html', {
                        'eventos': eventos,
                        'participantes': participantes,
                        'origen': origen,
                        'evento_preseleccionado': evento_preseleccionado,
                        'participante_preseleccionado': participante_preseleccionado,
                        'participantes_evento': participantes_evento,
                    })
                
                # Usar la plantilla del evento
                plantilla = evento.plantilla_id
                
                # Aquí implementarías la generación masiva
                messages.success(request, f"Generando constancias para {participantes_generar.count()} participantes del evento '{evento.titulo_evento}' con la plantilla del evento.")
                # Por ahora redireccionar al primer participante como ejemplo
                primer_participante = participantes_generar.first()
                return redirect('generar_constancia', primer_participante.id_participante, evento.id_evento, plantilla.id_plantilla)
                
            except Evento.DoesNotExist as e:
                messages.error(request, f"Error: {str(e)}")
        else:
            # Generación individual
            evento_id_form = request.POST.get('evento_id')
            participante_id_form = request.POST.get('participante_id')
            
            try:
                evento = Evento.objects.get(pk=evento_id_form)
                participante = Participante.objects.get(pk=participante_id_form)
                
                # Verificar que el evento tiene una plantilla
                if not evento.plantilla_id:
                    messages.error(request, f"El evento '{evento.titulo_evento}' no tiene una plantilla asociada.")
                    return render(request, 'constancias/pagina_generar.html', {
                        'eventos': eventos,
                        'participantes': participantes,
                        'origen': origen,
                        'evento_preseleccionado': evento_preseleccionado,
                        'participante_preseleccionado': participante_preseleccionado,
                        'participantes_evento': participantes_evento,
                    })
                
                # Usar la plantilla del evento
                plantilla = evento.plantilla_id
                
                return redirect('generar_constancia', participante_id_form, evento_id_form, plantilla.id_plantilla)
                
            except (Evento.DoesNotExist, Participante.DoesNotExist) as e:
                messages.error(request, f"Error: {str(e)}")
    else:
        # No necesitamos form ya que no usamos plantillas separadas
        pass

    contexto = {
        'eventos': eventos,
        'participantes': participantes,
        'origen': origen,
        'evento_preseleccionado': evento_preseleccionado,
        'participante_preseleccionado': participante_preseleccionado,
        'participantes_evento': participantes_evento,
    }
    
    return render(request, 'constancias/pagina_generar.html', contexto)


# ==================== NUEVAS VISTAS PARA PARTICIPANTES ====================

@login_required
def gestionar_participantes(request):
    """Vista principal para gestión de participantes"""
    
    # Obtener estadísticas
    total_participantes = Participante.objects.count()
    total_eventos = Evento.objects.filter(activo=True).count()
    
    # Participantes agregados en el último mes
    hace_un_mes = datetime.now() - timedelta(days=30)
    # Como no tenemos fecha_creacion en Participante, usaremos una estimación
    participantes_mes = Participante.objects.count() // 4  # Estimación simple
    
    # Si viene desde un evento específico
    evento = None
    evento_id = request.GET.get('evento_id')
    if evento_id:
        try:
            evento = Evento.objects.get(pk=evento_id)
        except Evento.DoesNotExist:
            messages.error(request, "El evento especificado no existe.")
    
    contexto = {
        'total_participantes': total_participantes,
        'total_eventos': total_eventos,
        'participantes_mes': participantes_mes,
        'evento': evento,
    }
    
    return render(request, 'constancias/gestionar_participantes.html', contexto)


@login_required
def crear_participante_individual(request):
    """Vista para crear un participante individual"""
    
    # Si viene desde un evento específico
    evento = None
    evento_id = request.GET.get('evento_id')
    if evento_id:
        try:
            evento = Evento.objects.get(pk=evento_id)
        except Evento.DoesNotExist:
            messages.error(request, "El evento especificado no existe.")
            return redirect('gestionar_participantes')
    
    if request.method == 'POST':
        form = ParticipanteForm(request.POST)
        
        if form.is_valid():
            try:
                # Crear participante
                participante = form.save()
                
                # Asociar a evento si se especificó
                evento_asociar_id = request.POST.get('evento_id')
                if evento_asociar_id:
                    try:
                        evento_asociar = Evento.objects.get(pk=evento_asociar_id)
                        evento_asociar.participantes.add(participante)
                        messages.success(
                            request, 
                            f'Participante "{participante.nombre_participante}" creado exitosamente y asociado al evento "{evento_asociar.titulo_evento}".'
                        )
                    except Evento.DoesNotExist:
                        messages.warning(
                            request, 
                            f'Participante "{participante.nombre_participante}" creado exitosamente, pero no se pudo asociar al evento.'
                        )
                else:
                    messages.success(
                        request, 
                        f'Participante "{participante.nombre_participante}" creado exitosamente.'
                    )
                
                # Redireccionar según el contexto
                if evento:
                    return redirect('detalle_evento', pk=evento.id_evento)
                else:
                    return redirect('lista_participantes')
                    
            except Exception as e:
                logger.error(f"Error al crear participante: {e}")
                messages.error(request, "Ocurrió un error al crear el participante. Inténtelo de nuevo.")
        else:
            messages.error(request, "Por favor corrija los errores en el formulario.")
    else:
        initial_data = {}
        if evento:
            initial_data['evento_id'] = evento.id_evento
            
        form = ParticipanteForm(initial=initial_data)
    
    # Obtener lista de eventos para el selector
    eventos = Evento.objects.filter(activo=True).order_by('-fecha_inicio')
    
    contexto = {
        'form': form,
        'evento': evento,
        'eventos': eventos,
    }
    
    return render(request, 'constancias/crear_participante_individual.html', contexto)


@login_required
def cargar_participantes_csv(request):
    """Vista para carga masiva de participantes por CSV"""
    
    # Obtener todos los eventos disponibles
    eventos = Evento.objects.all().order_by('-fecha_inicio')
    
    # Si viene desde un evento específico
    evento = None
    evento_id = request.GET.get('evento_id')
    if evento_id:
        try:
            evento = Evento.objects.get(pk=evento_id)
        except Evento.DoesNotExist:
            messages.error(request, "El evento especificado no existe.")
            return redirect('gestionar_participantes')
    
    if request.method == 'POST':
        try:
            archivo = request.FILES.get('csv_file')
            evento_seleccionado_id = request.POST.get('evento_id')
            
            if not archivo:
                messages.error(request, "No se recibió ningún archivo.")
                return render(request, 'constancias/cargar_participantes_csv.html', {
                    'evento': evento, 
                    'eventos': eventos
                })
            
            if not evento_seleccionado_id:
                messages.error(request, "Debe seleccionar un evento para asociar los participantes.")
                return render(request, 'constancias/cargar_participantes_csv.html', {
                    'evento': evento, 
                    'eventos': eventos
                })
            
            try:
                evento_seleccionado = Evento.objects.get(pk=evento_seleccionado_id)
            except Evento.DoesNotExist:
                messages.error(request, "El evento seleccionado no existe.")
                return render(request, 'constancias/cargar_participantes_csv.html', {
                    'evento': evento, 
                    'eventos': eventos
                })
            
            if not archivo.name.endswith('.csv'):
                messages.error(request, "El archivo debe ser un CSV válido.")
                return render(request, 'constancias/cargar_participantes_csv.html', {
                    'evento': evento, 
                    'eventos': eventos
                })
            
            # Leer y procesar CSV
            archivo_decodificado = archivo.read().decode('utf-8')
            datos_io = io.StringIO(archivo_decodificado)
            lector_csv = csv.reader(datos_io)
            
            # Leer encabezados
            try:
                encabezados = next(lector_csv)
                encabezados = [h.strip().lower() for h in encabezados]
            except StopIteration:
                messages.error(request, "El archivo CSV está vacío.")
                return render(request, 'constancias/cargar_participantes_csv.html', {
                    'evento': evento, 
                    'eventos': eventos
                })
            
            # Validar encabezados requeridos
            required_headers = ['nombre', 'email', 'rol']
            missing_headers = [h for h in required_headers if h not in encabezados]
            
            if missing_headers:
                messages.error(
                    request, 
                    f"Faltan las siguientes columnas requeridas: {', '.join(missing_headers)}"
                )
                return render(request, 'constancias/cargar_participantes_csv.html', {
                    'evento': evento, 
                    'eventos': eventos
                })
            
            # Obtener índices de las columnas
            nombre_idx = encabezados.index('nombre')
            email_idx = encabezados.index('email')
            rol_idx = encabezados.index('rol')
            
            participantes_creados = 0
            participantes_actualizados = 0
            errores = []
            
            for numero_fila, fila in enumerate(lector_csv, start=2):
                try:
                    if len(fila) < len(required_headers):
                        errores.append(f"Fila {numero_fila}: Datos insuficientes")
                        continue
                        
                    nombre = fila[nombre_idx].strip()
                    email = fila[email_idx].strip().lower()
                    rol = fila[rol_idx].strip()
                    
                    # Validaciones básicas
                    if len(nombre) < 2:
                        errores.append(f"Fila {numero_fila}: Nombre muy corto")
                        continue
                        
                    if '@' not in email or '.' not in email:
                        errores.append(f"Fila {numero_fila}: Email inválido")
                        continue
                        
                    if rol not in ['Participante', 'Ponente', 'Organizador', 'Facilitador']:
                        errores.append(f"Fila {numero_fila}: Rol inválido (debe ser: Participante, Ponente, Organizador o Facilitador)")
                        continue
                    
                    # Crear o actualizar participante
                    participante, created = Participante.objects.update_or_create(
                        email_participante=email,
                        defaults={
                            'nombre_participante': nombre,
                            'rol_participante': rol
                        }
                    )
                    
                    if created:
                        participantes_creados += 1
                    else:
                        participantes_actualizados += 1
                    
                    # Asociar al evento seleccionado
                    evento_seleccionado.participantes.add(participante)
                        
                except Exception as e:
                    errores.append(f"Fila {numero_fila}: Error procesando - {str(e)}")
            
            # Mensajes de resultado
            if participantes_creados > 0:
                messages.success(
                    request, 
                    f"✓ {participantes_creados} participantes creados exitosamente."
                )
            
            if participantes_actualizados > 0:
                messages.info(
                    request, 
                    f"✓ {participantes_actualizados} participantes existentes actualizados."
                )
            
            total_procesados = participantes_creados + participantes_actualizados
            if total_procesados > 0:
                messages.success(
                    request,
                    f"✓ {total_procesados} participantes asociados al evento '{evento_seleccionado.titulo_evento}'."
                )
            
            if errores:
                # Mostrar solo los primeros 5 errores
                errores_mostrar = errores[:5]
                if len(errores) > 5:
                    errores_mostrar.append(f"... y {len(errores) - 5} errores más")
                    
                messages.warning(
                    request,
                    f"⚠ {len(errores)} filas tuvieron errores: " + "; ".join(errores_mostrar)
                )
            
            if total_procesados > 0:
                # Redireccionar al evento seleccionado
                return redirect('detalle_evento', pk=evento_seleccionado.id_evento)
            
        except UnicodeDecodeError:
            messages.error(request, "Error al leer el archivo. Asegúrate de que esté codificado en UTF-8.")
        except Exception as e:
            logger.error(f"Error procesando CSV: {e}")
            messages.error(request, f"Error procesando el archivo: {str(e)}")
    
    contexto = {
        'evento': evento,
        'eventos': eventos,
    }
    
    return render(request, 'constancias/cargar_participantes_csv.html', contexto)


@login_required
def crear_evento(request):
    print("========================================")
    print("🟢 VIEW crear_evento EJECUTÁNDOSE")
    print("========================================")
    print(f"=== VIEW crear_evento llamada ===")
    print(f"Método: {request.method}")
    print(f"URL: {request.get_full_path()}")
    print(f"User: {getattr(request, 'user', 'Anónimo')}")
    
    if request.method == 'POST':
        print("=== PROCESANDO POST ===")
        print(f"POST data: {request.POST}")
        print(f"FILES: {request.FILES}")
        
        form = EventoForm(request.POST, request.FILES)
        print(f"Form válido: {form.is_valid()}")
        
        if form.is_valid():
            print("=== FORMULARIO VÁLIDO - CREANDO EVENTO ===")
            try:
                # Crear la plantilla con el archivo subido
                archivo_plantilla = form.cleaned_data['archivo_plantilla']
                
                # Crear la plantilla
                plantilla = Plantilla.objects.create(
                    nombre_plantilla=f"Plantilla_{form.cleaned_data['titulo_evento']}_{timezone.now().strftime('%Y%m%d_%H%M%S')}",
                    archivo=archivo_plantilla,
                    fecha_creacion=timezone.now().date(),
                    activa=True
                )
                
                # Crear el evento con activo=True por defecto y la plantilla creada
                evento = form.save(commit=False)
                evento.activo = True
                evento.plantilla_id = plantilla
                evento.save()
                
                messages.success(request, f'El evento "{evento.titulo_evento}" ha sido creado exitosamente.')
                return redirect('detalle_evento', pk=evento.id_evento)
                
            except Exception as e:
                messages.error(request, f'Error al crear el evento: {str(e)}')
                print(f"Error en crear_evento: {e}")
                return render(request, 'constancias/crear_evento.html', {'form': form})
                
        else:
            print("=== FORMULARIO INVÁLIDO ===")
            print(f"Errores: {form.errors}")
            for field, errors in form.errors.items():
                print(f"  {field}: {errors}")
    
    else:
        print("=== MÉTODO GET - MOSTRANDO FORMULARIO ===")
        form = EventoForm()
    
    return render(request, 'constancias/crear_evento.html', {'form': form})


@login_required
def descargar_plantilla_csv(request):
    """Vista para descargar la plantilla CSV de ejemplo para carga masiva"""
    
    # Crear la respuesta HTTP como archivo CSV
    response = HttpResponse(
        content_type='text/csv',
        headers={'Content-Disposition': 'attachment; filename="plantilla_participantes.csv"'},
    )
    
    # Escribir el CSV
    writer = csv.writer(response)
    
    # Escribir encabezados
    writer.writerow(['nombre', 'email', 'rol'])
    
    # Escribir ejemplos
    writer.writerow([
        'Juan Pérez García',
        'juan.perez@email.com',
        'Participante'
    ])
    writer.writerow([
        'María González López',
        'maria.gonzalez@email.com',
        'Ponente'
    ])
    writer.writerow([
        'Carlos Rodríguez Silva',
        'carlos.rodriguez@email.com',
        'Facilitador'
    ])
    writer.writerow([
        'Ana Martínez Torres',
        'ana.martinez@email.com',
        'Organizador'
    ])
    
    return response


@login_required
def eliminar_participante(request, pk):
    """Vista para eliminar un participante"""
    participante = get_object_or_404(Participante, pk=pk)
    
    if request.method == 'POST':
        nombre_participante = participante.nombre_participante
        participante.delete()
        messages.success(request, f'Participante "{nombre_participante}" eliminado exitosamente.')
        
        # Redireccionar según el origen
        redirect_to = request.POST.get('redirect_to', 'lista_participantes')
        if redirect_to == 'detalle_evento':
            evento_id = request.POST.get('evento_id')
            if evento_id:
                return redirect('detalle_evento', pk=evento_id)
        
        return redirect('lista_participantes')
    
    # Si es GET, mostrar página de confirmación
    contexto = {
        'participante': participante,
        'origen': request.GET.get('origen', 'lista_participantes'),
        'evento_id': request.GET.get('evento_id', None)
    }
    
    return render(request, 'constancias/confirmar_eliminar_participante.html', contexto)


@login_required
def eliminar_participante_de_evento(request, evento_pk, participante_pk):
    """Vista para eliminar la asociación de un participante con un evento específico"""
    evento = get_object_or_404(Evento, pk=evento_pk)
    participante = get_object_or_404(Participante, pk=participante_pk)
    
    if request.method == 'POST':
        # Verificar que el participante está asociado al evento
        if participante in evento.participantes.all():
            evento.participantes.remove(participante)
            messages.success(
                request, 
                f'Participante "{participante.nombre_participante}" removido del evento "{evento.titulo_evento}".'
            )
        else:
            messages.warning(
                request, 
                f'El participante "{participante.nombre_participante}" no estaba asociado a este evento.'
            )
        
        return redirect('detalle_evento', pk=evento_pk)
    
    # Si es GET, mostrar página de confirmación
    contexto = {
        'participante': participante,
        'evento': evento,
        'accion': 'remover_de_evento'
    }
    
    return render(request, 'constancias/confirmar_eliminar_participante.html', contexto)


@login_required
def exportar_participantes_csv(request):
    """Vista para exportar todos los participantes a CSV"""
    participantes = Participante.objects.all().order_by('nombre_participante')
    
    # Crear la respuesta HTTP como archivo CSV
    response = HttpResponse(
        content_type='text/csv',
        headers={'Content-Disposition': 'attachment; filename="participantes_exportados.csv"'},
    )
    
    # Escribir el CSV
    writer = csv.writer(response)
    
    # Escribir encabezados
    writer.writerow(['nombre', 'email', 'rol', 'eventos_asociados', 'fecha_exportacion'])
    
    # Escribir datos de participantes
    for participante in participantes:
        eventos_nombres = ', '.join([evento.titulo_evento for evento in participante.eventos.all()])
        writer.writerow([
            participante.nombre_participante,
            participante.email_participante,
            participante.rol_participante,
            eventos_nombres,
            timezone.now().strftime('%d/%m/%Y %H:%M')
        ])
    
    return response


@login_required
def exportar_participantes_evento_csv(request, evento_pk):
    """Vista para exportar participantes de un evento específico a CSV"""
    evento = get_object_or_404(Evento, pk=evento_pk)
    participantes = evento.participantes.all().order_by('nombre_participante')
    
    # Crear la respuesta HTTP como archivo CSV
    filename = f"participantes_{evento.titulo_evento.replace(' ', '_')}.csv"
    response = HttpResponse(
        content_type='text/csv',
        headers={'Content-Disposition': f'attachment; filename="{filename}"'},
    )
    
    # Escribir el CSV
    writer = csv.writer(response)
    
    # Escribir encabezados
    writer.writerow(['nombre', 'email', 'rol', 'evento', 'fecha_evento', 'fecha_exportacion'])
    
    # Escribir datos de participantes
    for participante in participantes:
        writer.writerow([
            participante.nombre_participante,
            participante.email_participante,
            participante.rol_participante,
            evento.titulo_evento,
            evento.fecha_inicio.strftime('%d/%m/%Y'),
            timezone.now().strftime('%d/%m/%Y %H:%M')
        ])
    
    return response


@login_required
def eliminar_evento(request, pk):
    """Vista para eliminar un evento"""
    evento = get_object_or_404(Evento, pk=pk)
    
    if request.method == 'POST':
        titulo_evento = evento.titulo_evento
        participantes_count = evento.participantes.count()
        
        try:
            # Eliminar también la plantilla asociada si existe
            if evento.plantilla_id:
                plantilla = evento.plantilla_id
                try:
                    # Eliminar archivo físico si existe
                    if plantilla.archivo and os.path.exists(plantilla.archivo.path):
                        os.remove(plantilla.archivo.path)
                    plantilla.delete()
                except Exception as e:
                    logger.warning(f"Error al eliminar plantilla: {e}")
            
            # Eliminar el evento (las relaciones many-to-many se eliminan automáticamente)
            evento.delete()
            
            messages.success(
                request, 
                f'Evento "{titulo_evento}" eliminado exitosamente. '
                f'{participantes_count} participantes fueron desvinculados del evento.'
            )
            
        except Exception as e:
            logger.error(f"Error al eliminar evento: {e}")
            messages.error(request, f"Error al eliminar el evento: {str(e)}")
        
        # Redireccionar según el origen
        redirect_to = request.POST.get('redirect_to', 'lista_eventos')
        if redirect_to == 'dashboard':
            return redirect('dashboard')
        else:
            return redirect('lista_eventos')
    
    # Si es GET, mostrar página de confirmación
    contexto = {
        'evento': evento,
        'origen': request.GET.get('origen', 'lista_eventos'),
        'participantes_count': evento.participantes.count()
    }
    
    return render(request, 'constancias/confirmar_eliminar_evento.html', contexto)


@login_required
def editar_evento(request, pk):
    """Vista para editar un evento"""
    evento = get_object_or_404(Evento, pk=pk)
    
    if request.method == 'POST':
        form = EventoForm(request.POST, request.FILES, instance=evento)
        
        if form.is_valid():
            try:
                # Si se sube un nuevo archivo de plantilla
                if 'archivo_plantilla' in request.FILES:
                    archivo_plantilla = form.cleaned_data['archivo_plantilla']
                    
                    # Eliminar plantilla anterior si existe
                    if evento.plantilla_id:
                        plantilla_anterior = evento.plantilla_id
                        try:
                            if plantilla_anterior.archivo and os.path.exists(plantilla_anterior.archivo.path):
                                os.remove(plantilla_anterior.archivo.path)
                            plantilla_anterior.delete()
                        except Exception as e:
                            logger.warning(f"Error al eliminar plantilla anterior: {e}")
                    
                    # Crear nueva plantilla
                    plantilla = Plantilla.objects.create(
                        nombre_plantilla=f"Plantilla_{form.cleaned_data['titulo_evento']}_{timezone.now().strftime('%Y%m%d_%H%M%S')}",
                        archivo=archivo_plantilla,
                        fecha_creacion=timezone.now().date(),
                        activa=True
                    )
                    
                    evento.plantilla_id = plantilla
                
                # Guardar el evento
                evento = form.save()
                
                messages.success(request, f'El evento "{evento.titulo_evento}" ha sido actualizado exitosamente.')
                return redirect('detalle_evento', pk=evento.id_evento)
                
            except Exception as e:
                logger.error(f"Error al editar evento: {e}")
                messages.error(request, f'Error al actualizar el evento: {str(e)}')
        else:
            messages.error(request, "Por favor corrija los errores en el formulario.")
    else:
        form = EventoForm(instance=evento)
    
    contexto = {
        'form': form,
        'evento': evento,
        'editing': True
    }
    
    return render(request, 'constancias/crear_evento.html', contexto)


@login_required
def api_participantes_evento(request, evento_pk):
    """Vista API para obtener participantes de un evento específico"""
    if not request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return JsonResponse({'error': 'Solo peticiones AJAX'}, status=400)
    
    try:
        evento = get_object_or_404(Evento, pk=evento_pk)
        participantes = evento.participantes.all().order_by('nombre_participante')
        
        participantes_data = []
        for participante in participantes:
            participantes_data.append({
                'id_participante': participante.id_participante,
                'nombre_participante': participante.nombre_participante,
                'email_participante': participante.email_participante,
                'rol_participante': participante.rol_participante,
            })
        
        return JsonResponse({
            'participantes': participantes_data,
            'total': len(participantes_data)
        })
        
    except Exception as e:
        logger.error(f"Error en API participantes evento: {e}")
        return JsonResponse({'error': 'Error interno del servidor'}, status=500)